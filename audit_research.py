import os
import openai
import web_scrapper
import prompts
import tools 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from dotenv import load_dotenv
import json

# Load API Key
load_dotenv()
api_key = os.environ.get("OPENAI_API_KEY")
client = openai.OpenAI(api_key=api_key) if api_key else None

def format_text_in_paragraph(paragraph, text):
    """Helper to apply bold formatting within a paragraph."""
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part: continue
        if i % 2 == 0:
            paragraph.add_run(part)
        else:
            paragraph.add_run(part).bold = True

# --- 1. Logic for Markdown Tables (Backup) ---
def create_word_table(doc, table_lines):
    """Converts text pipes | into a real Word Table."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    clean_rows = [r for r in rows if not all('-' in c for c in r)]
    if not clean_rows:
        return

    table = doc.add_table(rows=len(clean_rows), cols=len(clean_rows[0]))
    table.style = 'Table Grid' 

    # Populate cells
    for r_idx, row_data in enumerate(clean_rows):
        row_cells = table.rows[r_idx].cells
        for c_idx, cell_text in enumerate(row_data):
            paragraph = row_cells[c_idx].paragraphs[0]
            format_text_in_paragraph(paragraph, cell_text)
            
            # Apply shading and bolding for header row
            if r_idx == 0:
                paragraph.runs[0].bold = True
                # Set background color for header row (light gray)
                shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading_elm)

# --- 2. Logic to Parse Markdown to Docx ---
def parse_markdown_to_doc(doc, markdown_content, website_url, table_data=None):
    """Parses markdown text into a Word Document, handling headers, lists, and tables."""
    lines = markdown_content.split('\n')
    in_table = False
    table_lines = []
    
    # Track the location of the Website URL to insert it
    url_inserted = False

    for line in lines:
        stripped_line = line.strip()

        # Handle Tables
        if stripped_line.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped_line)
            continue
        elif in_table:
            # Table ended
            create_word_table(doc, table_lines)
            in_table = False
            table_lines = []
        
        # Handle Headers
        if stripped_line.startswith('###'):
            doc.add_heading(stripped_line.lstrip('# ').strip(), level=3)
        elif stripped_line.startswith('##'):
            doc.add_heading(stripped_line.lstrip('## ').strip(), level=2)
        elif stripped_line.startswith('#'):
            doc.add_heading(stripped_line.lstrip('# ').strip(), level=1)
        
        # Handle Lists
        elif stripped_line.startswith('* '):
            paragraph = doc.add_paragraph(style='List Bullet')
            format_text_in_paragraph(paragraph, stripped_line.lstrip('* ').strip())
        
        # Handle Paragraphs (and URL insertion)
        elif stripped_line:
            if stripped_line.startswith("## 1. Client Overview & Core Strategy") and not url_inserted:
                # Insert the URL right below the main section header
                doc.add_paragraph(f"Website: {website_url}", style='Intense Quote')
                url_inserted = True

            paragraph = doc.add_paragraph()
            format_text_in_paragraph(paragraph, stripped_line)


# --- 3. Main Audit Orchestration Function ---
def run_master_audit(openai_client, client_name: str, website_url: str, g_clients, output_folder_id: str):
    """
    Executes the full marketing audit pipeline: scraping, gathering data, 
    AI generation, document creation, and upload.
    """
    print(f"[DEBUG] Starting full audit for: {client_name} ({website_url})")

    # --- 1. Scrape Website Content ---
    print("\n[DEBUG] 🌐 Scrapping client website content...")
    client_text = web_scrapper.scrape_webpage(website_url)

    if client_text.startswith("Scrape failed"):
        print(f" ❌ Fatal: Failed to scrape client website. Scraper output: {client_text}")
        return None, None, None
    
    print(" ✅ Website content extracted successfully.")
    #### TESTING
    print(f"client_text: {client_text}")
    # --- 2. Technical & Pagespeed Analysis ---
    print("\n[DEBUG] ⚙️ Running Pagespeed/Technical Analysis...")
    pagespeed_scores = tools.get_pagespeed_insights(website_url) 
    
    if "Error" in pagespeed_scores:
        print(f" ❌ Fatal: Failed to get Pagespeed data. Output: {pagespeed_scores}")
        # Continue the audit but provide an error message to the LLM
        pagespeed_scores = f"Technical Audit Failed. Error: {pagespeed_scores}"
    else:
        print(" ✅ Pagespeed report generated successfully.")
    
    # --- 3. Gather SEO Snapshot (Serper) ---
    print("\n[DEBUG] 🔎 Gathering SEO snapshot...")
    seo_snapshot = tools.get_seo_snapshot(website_url, client_name)
    
    if seo_snapshot.startswith("Error"):
        print(f" ❌ Fatal: Failed to get SEO snapshot. Output: {seo_snapshot}")
        return None, None, None # Critical failure
    
    print(" ✅ SEO snapshot generated successfully.")
    
    # --- 4. Gather Competitor Data (Serper + AI Analysis) ---
    print("\n[DEBUG] 🤝 Finding and analyzing competitors...")
    
    # 4.1 Get Raw Competitors
    competitors_data = tools.get_competitors(client_name)
    
    if competitors_data.startswith("Error"):
        print(f" ❌ Fatal: Failed to find competitors. Output: {competitors_data}")
        return None, None, None # Critical failure
    
    print(f"  Found raw competitors for AI analysis: {competitors_data[:100]}...")

    # 4.2 Generate JSON Table for Competitor Comparison (Strict output)
    try:
        table_completion = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": prompts.SYSTEM_PROMPT_AUDIT},
                {"role": "user", "content": prompts.USER_PROMPT_COMPETITOR_JSON.format(competitor_data=competitors_data)}
            ]
        )
        table_json_str = table_completion.choices[0].message.content
        table_json = json.loads(table_json_str)
        print("  ✅ Competitor Comparison Table (JSON) Generated.")
    except Exception as e:
        print(f"  ❌ AI Table Generation Error: {e}")
        table_json = None # Fallback to no table if generation fails
        print("  ⚠️ Continuing without the comparison table.")

    # 4.3 Generate long-form summary (for main doc)
    # The competitor_data is used in the main prompt as the long-form analysis.

    # --- 5. Determine Business Type ---
    # This is a placeholder; in a real app, this might come from a CRM or an earlier step.
    business_type = "Digital Marketing & Full-Service Agency"
    print(f"\n[DEBUG] Business Type assumed: {business_type}")

    # --- 6. AI Master Document Generation (GPT-4) ---
    print("\n[DEBUG] 🧠 Generating Master Audit Document via AI...")
    
    master_prompt_data = {
        "client_name": client_name,
        "client_text": client_text,
        "seo_snapshot": seo_snapshot,
        "competitor_data": json.dumps(table_json, indent=2) if table_json else competitors_data, # Use table JSON or raw data
        "pagespeed_scores": pagespeed_scores,
        "business_type": business_type
    }

    try:
        completion = openai_client.chat.completions.create(
            model="gpt-4o-mini", # Use a capable model for this critical task
            messages=[
                {"role": "system", "content": prompts.SYSTEM_PROMPT_AUDIT},
                {"role": "user", "content": prompts.USER_PROMPT_MASTER_AUDIT.format(**master_prompt_data)}
            ]
        )
        master_document_content = completion.choices[0].message.content
        print("  ✅ Master Audit Document Generated.")
    except Exception as e:
        print(f"  ❌ AI Generation Error: {e}")
        return None, None, None

    # --- 7. Extract Summary and Video Prompt (from the generated document) ---
    
    # 7.1 Extract Website Summary (The main body of text)
    # Simple extraction: everything after the first header and before the last section
    try:
        # Find the content related to the core value prop/summary, which is usually at the start.
        summary_match = master_document_content.split("## 1. Client Overview & Core Strategy", 1)[1]
        website_summary = summary_match.split("##", 1)[0].strip()
        #### TESTING
        print(f"website summary: \n{website_summary}")
        print("  ✅ Extracted Website Summary.")
    except:
        # Fallback to the first 500 chars if structured extraction fails
        website_summary = master_document_content[:500] 
        print("  ⚠️ Falling back to simple summary extraction.")

    # 7.2 Extract Video Prompt Description (The final recommendation/call to action)
    # Check if the expected section header is present
    expected_video_section = "## 4. Video Strategy Recommendation"
    try:
        if expected_video_section in master_document_content:
            video_prompt_description = master_document_content.split(expected_video_section, 1)[1].strip()
            # Ensure we only take text up to the next potential section or end of document
            if "## " in video_prompt_description:
                 video_prompt_description = video_prompt_description.split("## ", 1)[0].strip()
            
            print("  ✅ Extracted Video Prompt Description.")
        else:
            # Fallback for structured extraction failure
            video_prompt_description = "A strategic video recommendation emphasizing digital growth and marketing excellence."
            print("  ⚠️ Falling back to generic video prompt description (Structured section not found).")

    except Exception as e:
        # Catch unexpected errors during split/extraction
        video_prompt_description = "A strategic video recommendation emphasizing digital growth and marketing excellence."
        print(f"  ⚠️ Falling back to generic video prompt description (Error during extraction: {e}).")

        #### TESTING
        print(f"video_prompt_description: \n{video_prompt_description}")
        

    # --- 8. Save & Upload ---
    filename = f"{client_name} - MASTER MARKETING AUDIT.docx"
    if not os.path.exists("temp_outputs"): os.makedirs("temp_outputs")
    local_path = os.path.join("temp_outputs", filename)

    try:
        doc = Document()
        doc.add_heading(f'MASTER MARKETING AUDIT: {client_name}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Parse content, handling the table
        parse_markdown_to_doc(doc, master_document_content, website_url, table_data=table_json) 
        
        doc.save(local_path)
        print(f"  ✅ Saved local doc: {local_path}")
    except Exception as e:
        print(f"  ❌ Doc Save Error: {e}")
        return None, None, None

    # Upload: FIX APPLIED HERE -> Changed 'upload_docx' to 'upload_file_to_drive'
    audit_link = g_clients.upload_file_to_drive(local_path, filename, output_folder_id)

    # Clean up local file
    try:
        os.remove(local_path)
    except Exception as e:
        print(f"  ⚠️ Warning: Could not remove local file {local_path}: {e}")
        
    # Return outputs needed for the next phase (video generation)
    return audit_link, website_summary, video_prompt_description